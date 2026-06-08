from docx import Document
from docx.shared import Pt, RGBColor
import re

def md_to_docx(md_path, docx_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Acumular filas de tabla
        if '|' in line and line.strip().startswith('|'):
            table_rows.append(line)
            i += 1
            continue

        # Procesar tabla acumulada
        if table_rows:
            rows = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
            if rows:
                cols = [c.strip() for c in rows[0].split('|') if c.strip()]
                t = doc.add_table(rows=len(rows), cols=len(cols))
                t.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    cells = [c.strip() for c in row.split('|') if c.strip() != '']
                    for ci in range(min(len(cols), len(t.rows[ri].cells))):
                        cell_text = cells[ci] if ci < len(cells) else ''
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                        t.rows[ri].cells[ci].text = cell_text
                        if ri == 0:
                            for run in t.rows[ri].cells[ci].paragraphs[0].runs:
                                run.bold = True
            table_rows = []
            doc.add_paragraph()

        # Linea vacia
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Bloque de codigo
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        # Encabezados
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 60)
        elif re.match(r'^[-*] ', line):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            doc.add_paragraph(text)

        i += 1

    doc.save(docx_path)
    print(f'Guardado: {docx_path}')

base = r'c:\Users\laura\OneDrive\Escritorio\Ingenieria de Software\IS_202610_E2'
md_to_docx(f'{base}\\Informe-Final-Postmortem.md', f'{base}\\Informe-Final-Postmortem.docx')
md_to_docx(f'{base}\\Manual-de-Usuario.md', f'{base}\\Manual-de-Usuario.docx')
